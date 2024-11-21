from bs4 import BeautifulSoup
import requests
import re
from io import BytesIO
from PIL import Image
import os
from datetime import datetime
from docx import Document

from docx.shared import Inches


def CreateDir(params):
    dir_name = params
    if not os.path.isdir(dir_name):
        os.makedirs(dir_name)


def SaveFunction(doc, paragraphs):
    counter = 1
    for paragraph in paragraphs:
        if isinstance(paragraph, str):
            doc.add_paragraph(paragraph)
        elif isinstance(paragraph, Image.Image):
            try:
                # Save the image to a file
                img_name = datetime.now().strftime('%Y%m%d%H%M%S') + str(counter) + ".jpg"
                img_path = os.path.join("Photos", img_name)
                paragraph.save(img_path)
                # Add the image to the document
                doc.add_picture(img_path, width=Inches(6))
                counter += 1
            except Exception as e:
                print("error", e)

    # create Word document
    doc_name = datetime.now().strftime('%Y%m%d%H%M%S') + ".docx"
    doc_path = os.path.join("DailyNews", doc_name)
    doc.save(doc_path)
    print(f"Document saved at {doc_path}")

    return doc


def GetWP():
    r = requests.get("https://www.wp.pl/")
    print("Status: ", r.status_code)
    print(r.url)

    soup = BeautifulSoup(r.text, "html.parser")
    linksa = soup.findAll("a",{"class"  : "sc-1xdae4w-0 kHGjGj sc-1nkodm-3 eiFKJZ sc-1k2mbc5-0 irgSwQ"})
    links2 = soup.findAll("a", {"class": "sc-1xdae4w-0 kHGjGj sc-1nkodm-3 kkaHYG sc-1y97t2y-1 iHQCPx"})
    links3 = soup.findAll("a", {"class": "sc-1xdae4w-0 kHGjGj sc-1nkodm-3 erDGhE sc-1k2mbc5-0 irgSwQ"})
    paragraphs = []

    for item in linksa:

        children = item.findChildren("div")
        child = children[len(children) - 1]
        stream = re.sub("<div[^>]+>", "", str(child))

        full_stream = ( str(stream) + "\n" + str(item.get("href")))

        paragraphs.append(full_stream)
        ##image
        children2 = item.findChildren("img")
        counter = 1
        for img in children2:
            img_download = requests.get(img.get('src'))

            try:
                img = Image.open(BytesIO(img_download.content))
                img_name = datetime.now().strftime('%Y%m%d%H%M%S') + str(counter) + ".jpg"
                img_path = os.path.join("Photos", img_name)
                img.save(img_path)
                paragraphs.append(Image.open(img_path))
                counter += 1
            except Exception as e:
                print("error", e)

    for item in links2:

        children = item.findChildren("div")
        child = children[len(children) - 1]
        stream = re.sub("<div[^>]+>", "", str(child))

        full_stream = ( str(stream) + "\n" + str(item.get("href")))
        paragraphs.append(full_stream)
        ##image
        children2 = item.findChildren("img")
        counter = 1
        for img in children2:
            img_download = requests.get(img.get('src'))

            try:
                img = Image.open(BytesIO(img_download.content))
                img_name = datetime.now().strftime('%Y%m%d%H%M%S') + str(counter) + ".jpg"
                img_path = os.path.join("Photos", img_name)
                img.save(img_path)
                paragraphs.append(Image.open(img_path))
                counter += 1
            except Exception as e:
                print("error", e)

    for item in links3:

        children = item.findChildren("div")
        child = children[len(children) - 1]
        stream = re.sub("<div[^>]+>", "", str(child))

        full_stream = ( str(stream) + "\n" + str(item.get("href")))
        paragraphs.append(full_stream)
        ##image
        children2 = item.findChildren("img")
        counter = 1
        for img in children2:
            img_download = requests.get(img.get('src'))

            try:
                img = Image.open(BytesIO(img_download.content))
                img_name = datetime.now().strftime('%Y%m%d%H%M%S') + str(counter) + ".jpg"
                img_path = os.path.join("Photos", img_name)
                img.save(img_path)
                paragraphs.append(Image.open(img_path))
                counter += 1
            except Exception as e:
                print("error", e)

    doc = Document()
    doc = SaveFunction(doc, paragraphs)


def GetTvn():


    r = requests.get("https://tvn24.pl/")
    print("Status: ", r.status_code)
    print(r.url)

    paragraphs = []
    soup1 = BeautifulSoup(r.text, "html.parser")
    linksa = soup1.findAll("div", {"class": "top-story-container"})

    for item in linksa:

        children = item.findChildren("a")
        for child in children:
            if 'title' in child.attrs:
                print(child['title'])
                print(child['href'])

                full_stream = ( str(child['title']) + "\n" + str(child['href']))
                paragraphs.append(full_stream)

                children3 = item.findChildren("img")
                counter = 1

                for img in children3:
                    img_download = requests.get(img.get('src'))
                    print(img_download.url)
                    try:
                        img = Image.open(BytesIO(img_download.content))
                        img_name = datetime.now().strftime('%Y%m%d%H%M%S') + str(counter) + ".jpg"
                        img_path = os.path.join("Photos", img_name)
                        img.save(img_path)
                        paragraphs.append(Image.open(img_path))
                        counter += 1
                    except Exception as e:
                        print("error", e)




    linksb = soup1.findAll("div", {"class": "teaser-wrapper"})

    for item in linksb:

        children2 = item.findChildren("a")
        for child in children2:
            if 'title' in child.attrs:
                prefix = "przejdź do "
                line_new = child['title'].removeprefix(prefix)
                print(line_new)
                print(child['href'])
                full_stream = ( str(line_new) + "\n" + str(child['href']))
                paragraphs.append(full_stream)

                children4 = child.findChildren("img")

                counter = 1
                for picture in children4:
                    try:
                        img_tag = picture.find("srcset")

                        if img_tag is not None and 'src' in img_tag.attrs:
                            img_src = img_tag['src']

                            print(img_src)
                            print(img_src)
                            try:
                                picture1 = Image.open(BytesIO(img_src.content))
                                img_name = datetime.now().strftime('%Y%m%d%H%M%S') + str(counter) + ".jpg"
                                img_path = os.path.join("Photos", img_name)
                                picture1.save(img_path)
                                paragraphs.append(Image.open(img_path))
                                counter += 1
                            except Exception as e:
                                print("error", e)
                    except Exception as e:
                        print("error", e)

    doc = Document()
    doc = SaveFunction(doc, paragraphs)


CreateDir("DailyNews")
CreateDir("Photos")
GetWP()
GetTvn()


